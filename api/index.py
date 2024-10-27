from flask import Flask, jsonify, request, send_file, after_this_request
from bs4 import BeautifulSoup
import requests
import openai
from datetime import datetime
import os
from supabase import create_client, Client
from cachetools import TTLCache
from typing import Dict, List
from fpdf import FPDF  # Make sure this import is at the top of your file
import os
from datetime import datetime
from docx import Document  # Add this import at the top
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = Flask(__name__)
supabase: Client = create_client(
    supabase_url='https://bfskhzxnjscoujggswqw.supabase.co',
    supabase_key='eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImJmc2toenhuanNjb3VqZ2dzd3F3Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3MzAwMDk5MzgsImV4cCI6MjA0NTU4NTkzOH0.aVNdwVJ7VSZ7Sjq4mBOx17hnzmH-AgxwwkIv9fr1r3s'
)
mistral_api_key = "FhMGkm2shvRoKS59clMyPcRjL2KBnlwn"

# Common App essay prompts for 2024-2025
COMMON_APP_QUESTIONS = [
    {
        "id": 1,
        "question": "Some students have a background, identity, interest, or talent that is so meaningful they believe their application would be incomplete without it. If this sounds like you, then please share your story.",
        "max_words": 650
    },
    {
        "id": 2,
        "question": "The lessons we take from obstacles we encounter can be fundamental to later success. Recount a time when you faced a challenge, setback, or failure. How did it affect you, and what did you learn from the experience?",
        "max_words": 650
    },
    {
        "id": 3,
        "question": "Reflect on a time when you questioned or challenged a belief or idea. What prompted your thinking? What was the outcome?",
        "max_words": 650
    },
    {
        "id": 4,
        "question": "Reflect on something that someone has done for you that has made you happy or thankful in a surprising way. How has this gratitude affected or motivated you?",
        "max_words": 650
    },
    {
        "id": 5,
        "question": "Discuss an accomplishment, event, or realization that sparked a period of personal growth and a new understanding of yourself or others.",
        "max_words": 650
    },
    {
        "id": 6,
        "question": "Describe a topic, idea, or concept you find so engaging that it makes you lose all track of time. Why does it captivate you? What or who do you turn to when you want to learn more?",
        "max_words": 650
    },
    {
        "id": 7,
        "question": "Share an essay on any topic of your choice. It can be one you've already written, one that responds to a different prompt, or one of your own design.",
        "max_words": 650
    }
]

UNIVERSITY_SPECIFIC_QUESTIONS = {
    "brown": [
        {
            "id": 1,
            "question": "What three words best describe you?",
            "max_words": 3
        },
        {
            "id": 2,
            "question": "What is your most meaningful extracurricular commitment, and what would you like us to know about it?",
            "max_words": 100
        },
        {
            "id": 3,
            "question": "If you could teach a class on any one thing, whether academic or otherwise, what would it be?",
            "max_words": 100
        },
        {
            "id": 4,
            "question": "In one sentence, Why Brown?",
            "max_words": 50
        }
    ],
    "dartmouth": [
        {
            "id": 1,
            "question": "As you seek admission to Class of 2029, what aspects of the college's academic program, community, and/or campus environment attract your interest? How is Dartmouth a good fit for you?",
            "max_words": 100
        }
    ],
    "duke": [
        {
            "id": 1,
            "question": "What is your sense of a university and a community, and why do you consider it a good match for you?",
            "max_words": 250
        }
    ]
}

SCHOLARSHIP_QUESTIONS = [
    {
        "id": 1,
        "question": "Describe how your leadership experience has shaped who you are today.",
        "max_words": 500
    },
    {
        "id": 2,
        "question": "What are your academic and career goals, and how will this scholarship help you achieve them?",
        "max_words": 500
    }
]

@app.route('/api/python/questions', methods=['GET'])
def get_application_questions():
    app_type = request.args.get('type', '')
    university = request.args.get('university', '').lower()
    
    if app_type == 'University':
        response_data = {
            "questions": COMMON_APP_QUESTIONS
        }
        
        # Add university-specific questions if university is specified
        if university and university in UNIVERSITY_SPECIFIC_QUESTIONS:
            response_data["supplemental"] = UNIVERSITY_SPECIFIC_QUESTIONS[university]
        
        return jsonify(response_data)
    
    elif app_type == 'Scholarship':
        return jsonify({"questions": SCHOLARSHIP_QUESTIONS})
    
    else:
        return jsonify({"error": "Invalid application type"}), 400
university_cache = TTLCache(maxsize=100, ttl=3600)

@app.route('/api/python/scrape-university', methods=['GET'])
def scrape_university_data():
    university = request.args.get('university')
    if not university:
        return jsonify({"error": "University parameter is required"}), 400

    # Check cache first
    if university in university_cache:
        return jsonify(university_cache[university])

    try:
        # Initialize scraping session with headers
        session = requests.Session()
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8'
        }

        # Format university name for search
        search_query = f"{university} university admissions requirements site:.edu"
        search_url = f"https://www.google.com/search?q={requests.utils.quote(search_query)}"

        # Get search results
        response = session.get(search_url, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')

        # Find first .edu link
        edu_links = [link.get('href') for link in soup.find_all('a') 
                    if 'href' in link.attrs and '.edu' in link.get('href')]
        
        if not edu_links:
            raise Exception("No university website found")

        # Scrape university page
        university_url = edu_links[0]
        university_response = session.get(university_url, headers=headers)
        university_soup = BeautifulSoup(university_response.text, 'html.parser')

        # Initialize data structure
        university_data = {
            "name": university,
            "admissions_criteria": [],
            "popular_majors": [],
            "application_tips": [],
            "requirements": []
        }

        # Extract GPA requirements
        gpa_patterns = ['GPA', 'Grade Point Average']
        for pattern in gpa_patterns:
            gpa_elements = university_soup.find_all(text=lambda text: pattern in text if text else False)
            for element in gpa_elements:
                text = element.strip()
                if any(char.isdigit() for char in text):
                    university_data["admissions_criteria"].append(text)

        # Extract SAT/ACT requirements
        test_patterns = ['SAT', 'ACT']
        for pattern in test_patterns:
            test_elements = university_soup.find_all(text=lambda text: pattern in text if text else False)
            for element in test_elements:
                text = element.strip()
                if any(char.isdigit() for char in text):
                    university_data["admissions_criteria"].append(text)

        # Extract majors
        major_keywords = ['major', 'program', 'degree']
        for keyword in major_keywords:
            major_elements = university_soup.find_all(['h2', 'h3', 'h4', 'p'], 
                text=lambda text: keyword in text.lower() if text else False)
            for element in major_elements:
                if len(element.text.strip()) > 5:  # Filter out too short texts
                    university_data["popular_majors"].append(element.text.strip())

        # Extract application tips
        tip_keywords = ['tip', 'advice', 'recommend', 'suggest']
        for keyword in tip_keywords:
            tip_elements = university_soup.find_all(['li', 'p'], 
                text=lambda text: keyword in text.lower() if text else False)
            for element in tip_elements:
                if len(element.text.strip()) > 10:  # Filter out too short texts
                    university_data["application_tips"].append(element.text.strip())

        # Extract requirements
        req_keywords = ['require', 'must', 'need']
        for keyword in req_keywords:
            req_elements = university_soup.find_all(['li', 'p'], 
                text=lambda text: keyword in text.lower() if text else False)
            for element in req_elements:
                if len(element.text.strip()) > 10:  # Filter out too short texts
                    university_data["requirements"].append(element.text.strip())

        # Remove duplicates and limit results
        for key in university_data:
            if isinstance(university_data[key], list):
                university_data[key] = list(dict.fromkeys(university_data[key]))[:10]

        # Add source URL
        university_cache[university] = university_data
        return jsonify(university_data)
    except Exception as e:
        print(f"Scraping error: {str(e)}")
        return jsonify({
            "error": "Failed to scrape university data",
            "details": str(e)
        }), 500

@app.route('/api/python/generate-answer', methods=['POST'])
def generate_answer():
    try:
        data = request.json
        if not data or 'question' not in data or 'user_profile' not in data:
            return jsonify({"error": "Missing required parameters"}), 400

        # Get user profile
        try:
            user_result = supabase.table('users').select('*').eq('userid', data['user_profile']['id']).execute()
            if not user_result.data or len(user_result.data) == 0:
                return jsonify({"error": "User profile not found"}), 404
            user_data = user_result.data[0]
        except Exception as db_error:
            print(f"Database error: {str(db_error)}")
            return jsonify({"error": "Failed to fetch user profile", "details": str(db_error)}), 500

        # Get university data if available
        university_data = None
        if data.get('university'):
            try:
                scrape_response = requests.get(
                    f"http://localhost:5328/api/python/scrape-university?university={data['university']}"
                )
                if scrape_response.status_code == 200:
                    university_data = scrape_response.json()
            except Exception as scrape_error:
                print(f"Error fetching university data: {str(scrape_error)}")

        # Construct prompt (same as before)
        prompt = f"""
        Student Profile:
        - Academic Background:
          • Grade Level: {user_data.get('grade', 'N/A')}
          • GPA: {user_data.get('gpa', 'N/A')}
          • Standardized Tests: {user_data.get('standardized_tests', 'N/A')}
          • Highest Coursework: {user_data.get('highest_coursework', 'N/A')}
          • Academic Awards: {user_data.get('academic_awards', 'N/A')}
        
        - Extracurricular Profile:
          • Activities: {user_data.get('extracurricular_activities', 'N/A')}
          • Community Service: {user_data.get('community_service', 'N/A')}
          • Hobbies: {user_data.get('hobbies', 'N/A')}
        
        - College Preferences:
          • Intended Majors: {user_data.get('intended_majors', 'N/A')}
          • Preferred College Type: {user_data.get('preferred_college_type', 'N/A')}
          • First Generation Student: {'Yes' if user_data.get('first_generation_student') else 'No'}
        """

        if university_data:
            prompt += f"""
            University Information for {university_data['name']}:
            - Admission Criteria:
              {chr(10).join([f"• {criterion}" for criterion in university_data['admissions_criteria']])}
            - Popular Majors:
              {chr(10).join([f"• {major}" for major in university_data['popular_majors']])}
            - Application Tips:
              {chr(10).join([f"• {tip}" for tip in university_data['application_tips']])}
            - Requirements:
              {chr(10).join([f"• {req}" for req in university_data['requirements']])}
            """

        prompt += f"""
        Application Question: {data['question']}
        Instructions: Write a compelling response that aligns with the university's values, addressing all aspects of the prompt, while highlighting the student's background. Avoid including any unnecessary details or references or headings, just the answer text is required. Make it sound unique and sound like i wrote it. Don't make it too long or complex."""

        # Mistral API request
        mistral_data = {
            "model": "open-mistral-nemo",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an experienced college application counselor helping students write compelling essays."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": 1500,
            "temperature": 0.85,
            "presence_penalty": 0.6,
            "frequency_penalty": 0.8
        }

        mistral_response = requests.post(
            "https://api.mistral.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {mistral_api_key}",
                "Content-Type": "application/json"
            },
            json=mistral_data
        )

        if mistral_response.status_code == 200:
            response_data = mistral_response.json()
            generated_answer = response_data['choices'][0]['message']['content']

            # Update application with new question and answer
            try:
                current_app = supabase.table('applications').select('questions, answers').eq('applicationid', data.get('applicationId')).execute()
                
                if current_app.data and len(current_app.data) > 0:
                    # Get current questions and answers
                    current_questions = current_app.data[0].get('questions', [])
                    current_answers = current_app.data[0].get('answers', [])
                    
                    # Update with new question and answer
                    supabase.table('applications').update({
                        'questions': current_questions + [data['question']],
                        'answers': current_answers + [generated_answer]
                    }).eq('applicationid', data.get('applicationId')).execute()
                else:
                    # Create new application if it doesn't exist
                    supabase.table('applications').insert({
                        'applicationid': data.get('applicationId'),
                        'title': data.get('university', ''),
                        'userid': data['user_profile']['id'],
                        'date': datetime.now().isoformat(),
                        'type': 'college',
                        'questions': [data['question']],
                        'answers': [generated_answer]
                    }).execute()

            except Exception as update_error:
                print(f"Error updating application: {str(update_error)}")
                # Continue even if update fails

            # Return the response in the format expected by the frontend
            return jsonify({
                "choices": [{
                    "message": {
                        "content": generated_answer
                    }
                }]
            })
        else:
            return jsonify({
                "error": "Failed to generate answer from Mistral",
                "details": mistral_response.text
            }), 500

    except Exception as e:
        print(f"Error generating answer: {str(e)}")
        return jsonify({
            "error": "Failed to generate answer",
            "details": str(e)
        }), 500
@app.route('/api/python/generate-docx', methods=['POST'])
def generate_docx():
    try:
        data = request.json
        questions = data.get('questions', [])
        answers = data.get('answers', [])
        title = data.get('title', 'Application Details')

        doc = Document()
        
        # Add title
        title_heading = doc.add_heading(title, level=1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content
        for i, (q, a) in enumerate(zip(questions, answers)):
            # Add question
            question_text = q['question'] if isinstance(q, dict) else str(q)
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f"Question {i+1}: {question_text}")
            question_run.bold = True
            question_run.font.size = Pt(12)

            # Add answer
            answer_para = doc.add_paragraph()
            if isinstance(a, list):  # If answer is formatted with bold sections
                for part in a:
                    if isinstance(part, dict):
                        run = answer_para.add_run(part['text'])
                        run.bold = part.get('bold', False)
                        run.font.size = Pt(12)
            else:
                answer_text = str(a) if a else "No answer provided"
                answer_run = answer_para.add_run(f"Answer: {answer_text}")
                answer_run.font.size = Pt(12)

            # Add spacing
            doc.add_paragraph()

        filename = f"application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc_path = os.path.join(os.getcwd(), filename)
        doc.save(doc_path)

        return send_file(
            doc_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='application.docx'
        )

    except Exception as e:
        print(f"DOCX Generation Error: {str(e)}")
        return jsonify({
            "error": "Failed to generate DOCX",
            "details": str(e)
        }), 500





if __name__ == '__main__':
    app.run(port=5328)

# @app.route('/api/python/generate-answer', methods=['POST'])
# def generate_answer():
#     try:
#         data = request.json
#         if not data or 'question' not in data or 'user_profile' not in data:
#             return jsonify({"error": "Missing required parameters"}), 400

#         # Get user profile
#         try:
#             user_result = supabase.table('users').select('*').eq('userid', data['user_profile']['id']).execute()
#             if not user_result.data or len(user_result.data) == 0:
#                 return jsonify({"error": "User profile not found"}), 404
#             user_data = user_result.data[0]
#         except Exception as db_error:
#             print(f"Database error: {str(db_error)}")
#             return jsonify({"error": "Failed to fetch user profile", "details": str(db_error)}), 500

#         # Get university data if available
#         university_data = None
#         if data.get('university'):
#             try:
#                 scrape_response = requests.get(
#                     f"http://localhost:5328/api/python/scrape-university?university={data['university']}"
#                 )
#                 if scrape_response.status_code == 200:
#                     university_data = scrape_response.json()
#             except Exception as scrape_error:
#                 print(f"Error fetching university data: {str(scrape_error)}")

#         # Construct enhanced prompt with university data
#         prompt = f"""
#         Student Profile:
#         - Academic Background:
#           • Grade Level: {user_data.get('grade', 'N/A')}
#           • GPA: {user_data.get('gpa', 'N/A')}
#           • Standardized Tests: {user_data.get('standardized_tests', 'N/A')}
#           • Highest Coursework: {user_data.get('highest_coursework', 'N/A')}
#           • Academic Awards: {user_data.get('academic_awards', 'N/A')}
        
#         - Extracurricular Profile:
#           • Activities: {user_data.get('extracurricular_activities', 'N/A')}
#           • Community Service: {user_data.get('community_service', 'N/A')}
#           • Hobbies: {user_data.get('hobbies', 'N/A')}
        
#         - College Preferences:
#           • Intended Majors: {user_data.get('intended_majors', 'N/A')}
#           • Preferred College Type: {user_data.get('preferred_college_type', 'N/A')}
#           • First Generation Student: {'Yes' if user_data.get('first_generation_student') else 'No'}
#         """

#         # Add university-specific information if available
#         if university_data:
#             prompt += f"""
            
#         University Information for {university_data['name']}:
#         - Admission Criteria:
#           {chr(10).join([f"• {criterion}" for criterion in university_data['admissions_criteria']])}
        
#         - Popular Majors:
#           {chr(10).join([f"• {major}" for major in university_data['popular_majors']])}
        
#         - Application Tips:
#           {chr(10).join([f"• {tip}" for tip in university_data['application_tips']])}
        
#         - Requirements:
#           {chr(10).join([f"• {req}" for req in university_data['requirements']])}
#         """

#         prompt += f"""
        
#         Application Question: {data['question']}

#         Instructions for Response:
#         1. Write a compelling response that aligns with the university's values and requirements
#         2. Use specific examples from the student's background that match the university's criteria
#         3. Reference relevant university-specific information when applicable
#         4. Maintain a natural, conversational tone while being academically appropriate
#         5. Address all aspects of the prompt thoroughly
#         6. Stay within word limits while being detailed and specific
#         7. Highlight experiences that align with the university's popular programs and requirements
#         8. If the student is a first-generation student, incorporate that perspective thoughtfully
#         9. Connect the student's experiences to their intended major and the university's offerings
#         """

#         # Generate response using OpenAI
#         client = openai.Client(api_key=openai.api_key)
#         response = client.chat.completions.create(
#             model="gpt-4o-mini",
#             messages=[
#                 {
#                     "role": "system", 
#                     "content": "You are an experienced college application counselor who helps students write compelling, university-specific application essays."
#                 },
#                 {"role": "user", "content": prompt}
#             ],
#             max_tokens=1500,
#             temperature=0.85,
#             presence_penalty=0.6,
#             frequency_penalty=0.8
#         )

#         generated_answer = response.choices[0].message.content

#         # Update application with new answer (same as before)
#         try:
#             current_app = supabase.table('applications').select('questions_answers').eq('applicationid', data.get('applicationId')).execute()
#             if current_app.data and len(current_app.data) > 0:
#                 current_qa = current_app.data[0].get('questions_answers', {'questions': [], 'answers': []})
                
#                 supabase.table('applications').update({
#                     'questions_answers': {
#                         'questions': current_qa.get('questions', []) + [data['question']],
#                         'answers': current_qa.get('answers', []) + [generated_answer]
#                     }
#                 }).eq('applicationid', data.get('applicationId')).execute()
#         except Exception as update_error:
#             print(f"Error updating application: {str(update_error)}")

#         return jsonify({
#             "answer": generated_answer,
#             "timestamp": datetime.now().isoformat(),
#             "status": "success"
#         })

#     except Exception as e:
#         print(f"Error generating answer: {str(e)}")
#         return jsonify({
#             "error": "Failed to generate answer",
#             "details": str(e)
#         }), 500


    
    
#  fetch common app questions an pass it to frontend for answering

# fetch scholarship application form questions and pass it to frontend for answering

# scrape through university websites, articles and accepted student's application from previous admissions along wiht user's profile to develop context for the llm then ask specific questions requested as part of the application that the user is trying to fill out.  

# ai api with small prompts for that it is a counsellor bot with some knowledge about user's info